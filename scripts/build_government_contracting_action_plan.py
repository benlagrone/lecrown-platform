from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "doc"
DOCX_PATH = OUTPUT / "lecrown-government-contracting-action-plan-2026-07-28.docx"

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
GOLD = "C79A2B"
LIGHT_GOLD = "F7EBC8"
LIGHT_GRAY = "F2F2F2"
DARK = RGBColor(31, 41, 55)
WHITE = RGBColor(255, 255, 255)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="D9E2F3", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_status_callout(doc, title, body, fill=LIGHT_GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    shade(cell, fill)
    set_cell_margins(cell, 110, 130, 110, 130)
    p = cell.paragraphs[0]
    run = p.add_run(title + ": ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_action_table(doc, rows):
    headers = ["ID / priority", "Action", "Owner / due", "State", "Definition of done"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.75, 2.35, 1.05, 1.05, 2.25]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        values = [
            f"{row[0]}\n{row[1]}",
            row[2],
            f"{row[3]}\n{row[4]}",
            row[5],
            row[6],
        ]
        for col, value in enumerate(values):
            cell = cells[col]
            cell.width = Inches(widths[col])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if idx % 2:
                shade(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.size = Pt(7.5)
            if col == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(BLUE)
        prevent_row_split(table.rows[-1])
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        shade(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx, value in enumerate(row):
            cell = cells[cidx]
            cell.width = Inches(widths[cidx])
            set_cell_margins(cell)
            if ridx % 2:
                shade(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(8)
            if cidx == 0:
                run.bold = True
        prevent_row_split(table.rows[-1])
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
styles["Normal"].font.color.rgb = DARK
styles["Normal"].paragraph_format.space_after = Pt(5)
styles["Normal"].paragraph_format.line_spacing = 1.05
for name, size, color in (
    ("Title", 24, NAVY),
    ("Heading 1", 16, NAVY),
    ("Heading 2", 12, BLUE),
    ("Heading 3", 10, NAVY),
):
    style = styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.space_before = Pt(8)
    style.paragraph_format.space_after = Pt(4)

header = section.header
header_p = header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_run = header_p.add_run("LECROWN DEVELOPMENT | GOVERNMENT CONTRACTING")
header_run.font.name = "Aptos"
header_run.font.size = Pt(7.5)
header_run.font.bold = True
header_run.font.color.rgb = RGBColor.from_string(NAVY)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run("Internal working plan | Prepared 2026-07-28")
footer_run.font.name = "Aptos"
footer_run.font.size = Pt(7.5)
footer_run.font.color.rgb = RGBColor(89, 89, 89)

title = doc.add_paragraph()
title.style = "Title"
title.paragraph_format.space_before = Pt(14)
title.paragraph_format.space_after = Pt(2)
title.add_run("Government Contracting Action Plan")
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run("August 2026 execution plan | APEX follow-up milestone: August 10")
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string(GOLD)

add_status_callout(
    doc,
    "Primary objective",
    "Enter the August 10 meeting with one verified legal-entity profile, an advisor-ready capability statement, a controlled City OBO recovery plan, and a ranked prime/local-buyer pipeline.",
    LIGHT_BLUE,
)

add_heading(doc, "What is already complete", 1)
for text in (
    "Revised one-page capability statement sent to Marina on July 16 and audited against the APEX checklist.",
    "Texas CMBL publicly verified through April 2, 2027 with 23 NIGP codes and all 25 highway districts.",
    "Houston METRO SBE certification portal-verified from July 15, 2026 through July 15, 2029.",
    "APEX subcontracting reading, procurement-link review, SBA prime screening, and six-prime target list completed.",
):
    add_bullet(doc, text)

add_heading(doc, "Critical blockers", 1)
for text in (
    "Legal names are not normalized across LeCrown Development, LeCrown Development Corporation, and Metro LeCrown.",
    "SAM, UEI, CAGE, and SBA public-profile status are not verified in the tracker.",
    "NIGP 920-37 appears on the capability statement but not in the verified CMBL profile.",
    "City OBO application 6369689 is 52 percent complete with 0 of 13 mandatory documents attached and a September 12 deletion date.",
):
    add_bullet(doc, text)

add_heading(doc, "Operating rules", 1)
for text in (
    "One authoritative entity anchors all registrations, certifications, collateral, and supplier profiles.",
    "Prepared, submitted, sent, and publicly verified are separate statuses.",
    "Do not add UEI, CAGE, SAM status, or certifications until verified for the same entity.",
    "Do not represent key-personnel work as company prime-contract past performance.",
    "Do not submit portals, upload sensitive files, or send outreach without the required review gate.",
):
    add_number(doc, text)

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "Phase 1 - Stabilize identity and evidence", 1)
doc.add_paragraph("Target: July 28 through August 2")
phase1 = [
    ("A01", "CRITICAL", "Establish authoritative legal entity and map it to SAM, CMBL, METRO, City OBO, website, and collateral.", "Benjamin/Jie", "Jul 29", "Not started", "One signed-off entity map with evidence links."),
    ("A02", "CRITICAL", "Capture current SAM status, UEI, CAGE, renewal date, and public entity name.", "Benjamin", "Jul 30", "Not verified", "Tracker updated without sensitive banking or tax values."),
    ("A03", "HIGH", "Package current CMBL public proof and verified 23-code NIGP list.", "Benjamin", "Jul 29", "Evidence open", "Vendor ID, status, expiration, category, codes, and source retained."),
    ("A04", "HIGH", "Save METRO SBE proof and reconcile the certificate legal name.", "Benjamin/Jie", "Aug 2", "Evidence open", "Certificate linked and renewal date tracked."),
    ("A05", "HIGH", "Verify or create the project-scoped document-portal record.", "Benjamin", "Aug 2", "Not verified", "Project ID and first approved document record captured."),
    ("A06", "CRITICAL", "Inventory all 13 City OBO mandatory documents.", "Benjamin/Jie", "Jul 31", "In progress", "Each item has owner, source, storage, readiness, and upload decision."),
    ("A07", "CRITICAL", "Complete six City OBO form sections without signing or submitting.", "Benjamin/Jie", "Jul 31", "In progress", "All required sections complete; no submission performed."),
    ("A08", "HIGH", "Locate or create City Strategic Purchasing registration proof.", "Benjamin/Jie", "Aug 1", "Not verified", "Acceptable proof saved for the OBO packet."),
]
add_action_table(doc, phase1)
add_status_callout(
    doc,
    "Gate 1",
    "Do not release revised collateral, register with federal primes, or market METRO SBE until A01 and A02 are complete and A04 is reconciled.",
)

add_heading(doc, "Phase 2 - Advisor-ready package", 1)
doc.add_paragraph("Target: August 3 through August 9")
phase2 = [
    ("A09", "CRITICAL", "Correct capability-statement company data.", "Benjamin", "Aug 3", "Staged", "Legal name normalized; NIGP mismatch resolved; only verified company data included."),
    ("A10", "HIGH", "Align website and capability-statement claims.", "Benjamin", "Aug 3", "Not started", "Names, services, credentials, outcomes, and evidence boundaries agree."),
    ("A11", "HIGH", "Export and visually inspect advisor-review PDF.", "Benjamin", "Aug 4", "Not started", "Readable one-page PDF under 1 MB; July 16 version retained."),
    ("A12", "HIGH", "Confirm which version Marina is reviewing.", "Benjamin", "Aug 4", "Awaiting contact", "Marina preference recorded; send occurs only after review."),
    ("A13", "MEDIUM", "Verify SBA Dynamic Small Business Search profile.", "Benjamin", "Aug 5", "Blocked by A02", "Public entity, codes, keywords, and verified certifications are correct."),
    ("A14", "MEDIUM", "Run MySBA WOSB/EDWOSB, 8(a), and HUBZone checks.", "Benjamin/Jie", "Aug 7", "Queued", "Outcomes recorded and one lane selected or intentionally deferred."),
    ("A15", "MEDIUM", "Create UH/UT and Texas ESBD saved-search routine.", "Benjamin", "Aug 7", "Queued", "Codes, terms, owner, frequency, and evidence path documented."),
    ("A16", "HIGH", "Prepare August 10 APEX meeting brief.", "Benjamin", "Aug 9", "Not started", "Brief covers progress, blockers, six primes, OBO risk, and advisor questions."),
]
add_action_table(doc, phase2)

add_heading(doc, "August 10 decision agenda", 2)
for number, text in enumerate((
    "Confirm which capability statement Marina is reviewing.",
    "Confirm the legal entity and certifications for federal-prime marketing.",
    "Select the first market lane: federal primes, Houston local agencies, or one named opportunity.",
    "Resolve whether NIGP 920-37 should be added to CMBL or removed from collateral.",
    "Select the first two prime targets.",
    "Ask for the July 16 recording/transcript or confirmation that none is available.",
), start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"{number}. ").bold = True
    p.add_run(text)

add_heading(doc, "Phase 3 - Approved prime registrations and outreach", 1)
doc.add_paragraph("Target: August 11 through August 21")
phase3 = [
    ("A17", "HIGH", "Register with first two Marina-approved prime programs.", "Benjamin", "Aug 13", "Approval required", "Profile IDs, dates, capabilities, uploads, and proof recorded."),
    ("A18", "HIGH", "Register with next two qualified primes.", "Benjamin", "Aug 17", "Blocked by A17", "Four total federal-prime profiles are verified."),
    ("A19", "HIGH", "Build one opportunity hypothesis for each registered prime.", "Benjamin", "Aug 18", "Not started", "Each has an agency/vehicle, problem, LeCrown fit, and routing path."),
    ("A20", "HIGH", "Draft opportunity-specific outreach.", "Benjamin", "Aug 19", "Not started", "Four drafts use verified data and accurate experience labels."),
    ("A21", "HIGH", "Review and send first outreach wave.", "Benjamin", "Aug 20", "Send approval", "Sent proof retained and follow-up dates scheduled."),
    ("A22", "MEDIUM", "Resolve Dell Federal subcontracting route.", "Benjamin", "Aug 18", "Research open", "Authoritative federal route recorded or target deprioritized."),
    ("A23", "MEDIUM", "Resolve KBR SBLO/supplier route.", "Benjamin", "Aug 18", "Research open", "Authoritative route recorded or target deprioritized."),
]
add_action_table(doc, phase3)

doc.add_page_break()
add_heading(doc, "Phase 4 - Local buyer readiness", 1)
doc.add_paragraph("Target: August 11 through August 28")
phase4 = [
    ("A24", "CRITICAL", "Finish City OBO document assembly and internal review.", "Benjamin/Jie", "Aug 14", "In progress", "All 13 items ready or exceptions have dated resolution plans."),
    ("A25", "CRITICAL", "Review completed City OBO application before submission.", "Benjamin/Jie", "Aug 17", "Approval required", "Facts, control answers, attachments, signatures, and claims pass review."),
    ("A26", "CRITICAL", "Submit City OBO only after explicit approval.", "Benjamin/Jie", "Aug 18", "Not authorized", "Confirmation and submitted package retained; status rechecked."),
    ("A27", "HIGH", "Verify City BeaconBid and Strategic Purchasing profiles.", "Benjamin", "Aug 19", "Not verified", "Accounts, commodity settings, and notifications evidenced."),
    ("A28", "MEDIUM", "Verify Harris County Bonfire and UH/ESBD monitoring.", "Benjamin", "Aug 21", "Not verified", "Registration/search evidence and alerts saved."),
    ("A29", "MEDIUM", "Work remaining education buyer portals in ranked order.", "Benjamin", "Aug 28", "Queued", "Each portal verified, deferred, or marked not applicable."),
]
add_action_table(doc, phase4)

add_heading(doc, "Weekly operating cadence", 1)
cadence_rows = [
    ("Monday", "Pipeline scan", "SAM, SBA/DSBS, USAspending, Texas ESBD, METRO, City of Houston, and priority buyers."),
    ("Tuesday", "Registrations and evidence", "Complete one approved profile or evidence task and update the tracker immediately."),
    ("Wednesday", "Prime research and outreach", "Advance two primes and tie each message to a named opportunity, agency, or vehicle."),
    ("Thursday", "Certification packet", "Advance City OBO and the selected federal-certification lane; keep sensitive data private."),
    ("Friday", "Follow-up and scorecard", "Follow up after five business days, reconcile status to proof, and set next week's top three."),
]
add_simple_table(doc, ["Day", "Focus", "Minimum output"], cadence_rows, [0.8, 1.7, 4.6])

doc.add_page_break()
add_heading(doc, "August scorecard", 1)
score_rows = [
    ("Authoritative entity profiles reconciled", "1"),
    ("Advisor-ready capability statements", "1 master plus 2 targeted variants"),
    ("Prime supplier profiles verified", "4"),
    ("Opportunity hypotheses completed", "4"),
    ("Targeted outreach messages sent and evidenced", "4"),
    ("Qualified prime conversations booked", "2"),
    ("Local buyer registrations verified", "4"),
    ("City OBO mandatory documents ready", "13 of 13"),
    ("City OBO application sections complete", "100 percent"),
    ("Unsupported completion claims", "0"),
]
add_simple_table(doc, ["Measure", "August target"], score_rows, [5.5, 1.6])

add_heading(doc, "Approval gates", 1)
for text in (
    "CMBL code changes that can trigger submission or payment.",
    "City OBO signature or submission.",
    "Upload of sensitive ownership, tax, banking, or corporate records.",
    "Revised capability statement sent to Marina.",
    "Prime or buyer profile submission.",
    "Prime-contractor outreach.",
    "Public website publication.",
):
    add_bullet(doc, text)

add_status_callout(
    doc,
    "Evidence standard",
    "A local draft, completed form, submitted profile, sent message, portal confirmation, and public verification are separate claims. Record the strongest evidence actually obtained.",
    LIGHT_BLUE,
)

add_heading(doc, "Source files", 1)
source_text = "; ".join((
    "docs/apex-subcontracting-homework-2026-08-10.md",
    "data/prime-subcontracting-targets.csv",
    "data/certification-work-queue.csv",
    "data/company-certifications.csv",
    "data/buyer-certification-requirements.csv",
    "docs/certification-tracker.md",
    "/Users/benjaminlagrone/Documents/projects/lecrowndev/site/output/pdf/lecrown-development-capability-statement.pdf",
))
source_p = doc.add_paragraph()
source_p.paragraph_format.space_after = Pt(0)
source_run = source_p.add_run(source_text)
source_run.font.size = Pt(7.5)

OUTPUT.mkdir(parents=True, exist_ok=True)
doc.save(DOCX_PATH)
print(DOCX_PATH)
